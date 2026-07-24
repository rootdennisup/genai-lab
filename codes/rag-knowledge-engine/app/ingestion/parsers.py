"""四类文档解析器。

解析阶段只负责把不同文件格式统一成 ``ParsedBlock``，不在这里做 Chunk，
这是“格式知识”和“RAG 策略”之间的重要边界。
"""

from __future__ import annotations

import re
from pathlib import Path

from app.domain.models import ParsedBlock


class ParseError(ValueError):
    """文件无法安全、可靠地解析时抛出。"""


def parse_document(path: Path, file_type: str) -> list[ParsedBlock]:
    parsers = {"md": _parse_markdown, "pdf": _parse_pdf, "docx": _parse_docx, "xlsx": _parse_xlsx}
    try:
        return parsers[file_type](path)
    except KeyError as exc:
        raise ParseError(f"不支持的文件类型：{file_type}") from exc
    except ParseError:
        raise
    except Exception as exc:
        # 不把第三方库堆栈直接暴露给 API 使用者，但保留清晰的失败类别。
        raise ParseError(f"{file_type.upper()} 文档解析失败：{exc}") from exc


def _clean(text: str) -> str:
    """保守清洗：统一空白但保留标点、列表和专有名词。"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _parse_markdown(path: Path) -> list[ParsedBlock]:
    # 解析 Markdown 需要安装 markdown-it-py。
    lines = path.read_text(encoding="utf-8-sig").splitlines()
    blocks: list[ParsedBlock] = []
    headings: list[str] = []
    buffer: list[str] = []
    start_line = 1

    def flush(end_line: int) -> None:
        nonlocal buffer, start_line
        text = _clean("\n".join(buffer))
        if text:
            blocks.append(ParsedBlock(text=text, title=headings[-1] if headings else path.stem,
                                      section_path=headings.copy(),
                                      locator={"line_start": start_line, "line_end": end_line}))
        buffer = []

    for number, line in enumerate(lines, 1):
        match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if match:
            flush(number - 1)
            level, title = len(match.group(1)), match.group(2).strip()
            headings[level - 1:] = [title]
            start_line = number + 1
        else:
            if not buffer:
                start_line = number
            buffer.append(line)
    flush(len(lines))
    return blocks


def _parse_pdf(path: Path) -> list[ParsedBlock]:
    # 解析 PDF 需要安装 pypdf，MVP 暂不支持 OCR。
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ParseError("解析 PDF 需要安装 pypdf") from exc
    blocks = []
    for page_number, page in enumerate(PdfReader(path).pages, 1):
        text = _clean(page.extract_text() or "")
        if text:
            blocks.append(ParsedBlock(text=text, title=path.stem,
                                      locator={"page_start": page_number, "page_end": page_number}))
    if not blocks:
        raise ParseError("PDF 未提取到文本，可能是扫描件；MVP 暂不支持 OCR")
    return blocks


def _parse_docx(path: Path) -> list[ParsedBlock]:
    # 解析 Word 需要安装 python-docx。
    try:
        from docx import Document
    except ImportError as exc:
        raise ParseError("解析 Word 需要安装 python-docx") from exc
    document = Document(path)
    headings: list[str] = []
    blocks: list[ParsedBlock] = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = _clean(paragraph.text)
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        match = re.search(r"Heading\s*(\d+)|标题\s*(\d+)", style, re.I)
        if match:
            level = int(match.group(1) or match.group(2))
            headings[level - 1:] = [text]
            continue
        blocks.append(ParsedBlock(text=text, title=headings[-1] if headings else path.stem,
                                  section_path=headings.copy(), locator={"paragraph": index}))
    # 表格按“字段名：值”展开，让每一行在脱离原表后仍保留列语义。
    for table_index, table in enumerate(document.tables, 1):
        rows = [[_clean(cell.text) for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        headers = rows[0]
        for row_index, values in enumerate(rows[1:], 2):
            text = "\n".join(f"{headers[i] or f'列{i + 1}'}：{value}" for i, value in enumerate(values) if value)
            if text:
                blocks.append(ParsedBlock(text=text, title=f"表格 {table_index}",
                                          locator={"table": table_index, "row_start": row_index,
                                                   "row_end": row_index}))
    return blocks


def _parse_xlsx(path: Path) -> list[ParsedBlock]:
    # 解析 Excel 需要安装 openpyxl。
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ParseError("解析 Excel 需要安装 openpyxl") from exc
    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[ParsedBlock] = []
    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(value).strip() if value is not None else f"列{i + 1}" for i, value in enumerate(rows[0])]
        for row_number, row in enumerate(rows[1:], 2):
            fields = [f"{headers[i]}：{value}" for i, value in enumerate(row) if value is not None]
            if fields:
                blocks.append(ParsedBlock(text="\n".join(fields), title=sheet.title,
                                          section_path=[sheet.title],
                                          locator={"sheet_name": sheet.title, "row_start": row_number,
                                                   "row_end": row_number}))
    return blocks

